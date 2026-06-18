"""PreMortem report exporters: JSON, PDF and Word."""
from __future__ import annotations

import io
import json

from ..models import PreMortemReport


def to_json_bytes(report: PreMortemReport) -> bytes:
    return json.dumps(report.model_dump(), indent=2, default=str).encode("utf-8")


def to_pdf_bytes(report: PreMortemReport) -> bytes:
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import mm
    from reportlab.platypus import (
        ListFlowable,
        ListItem,
        Paragraph,
        SimpleDocTemplate,
        Spacer,
        Table,
        TableStyle,
    )

    buf = io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=A4, title="PreMortem Report")
    styles = getSampleStyleSheet()
    h1 = ParagraphStyle("h1", parent=styles["Title"], textColor=colors.HexColor("#0f172a"))
    h2 = ParagraphStyle(
        "h2", parent=styles["Heading2"], textColor=colors.HexColor("#1e3a8a")
    )
    body = styles["BodyText"]

    decision_color = {
        "GO": "#16a34a",
        "GO WITH CONDITIONS": "#d97706",
        "NO-GO": "#dc2626",
    }.get(report.recommended_decision.value, "#dc2626")

    story = []
    story.append(Paragraph("PreMortem AI - Procurement Risk Report", h1))
    story.append(Spacer(1, 4 * mm))

    summary = [
        ["Procurement", report.procurement_name],
        ["Equipment", report.equipment_type],
        ["Contract Value", f"Rs {report.contract_value_cr:.2f} Cr"],
        ["Overall Risk Score", f"{report.overall_risk_score:.0f}/100"],
        ["Failure Probability", f"{report.failure_probability_pct:.0f}%"],
        ["Confidence", f"{report.confidence_pct:.0f}%"],
        ["Predicted Delay", f"{report.predicted_delay_months:.0f} months"],
        ["Projected Loss", f"Rs {report.projected_financial_loss_cr:.2f} Cr"],
    ]
    t = Table(summary, colWidths=[55 * mm, 110 * mm])
    t.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (0, -1), colors.HexColor("#e2e8f0")),
                ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#cbd5e1")),
                ("FONTSIZE", (0, 0), (-1, -1), 10),
                ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
            ]
        )
    )
    story.append(t)
    story.append(Spacer(1, 5 * mm))

    story.append(
        Paragraph(
            f'<font color="{decision_color}"><b>Recommended Decision: '
            f"{report.recommended_decision.value}</b></font>",
            h2,
        )
    )
    story.append(Spacer(1, 3 * mm))

    story.append(Paragraph("Predicted Failure Mode", h2))
    story.append(Paragraph(report.predicted_failure_mode, body))
    story.append(Spacer(1, 3 * mm))

    def bullet_list(items):
        return ListFlowable(
            [ListItem(Paragraph(str(i), body)) for i in items],
            bulletType="bullet",
        )

    story.append(Paragraph("Supporting Evidence", h2))
    story.append(bullet_list(report.supporting_evidence))
    story.append(Spacer(1, 3 * mm))

    story.append(Paragraph("Predicted Outcomes", h2))
    story.append(bullet_list(report.predicted_outcomes))
    story.append(Spacer(1, 3 * mm))

    if report.conditions:
        story.append(Paragraph("Conditions for Approval", h2))
        story.append(bullet_list(report.conditions))
        story.append(Spacer(1, 3 * mm))

    story.append(Paragraph("Agent Findings", h2))
    for r in report.agent_results:
        story.append(
            Paragraph(
                f"<b>{r.agent}</b> - Risk {r.risk_score:.0f}/100 ({r.risk_level.value})",
                body,
            )
        )
        if r.findings:
            story.append(bullet_list(r.findings))
        story.append(Spacer(1, 2 * mm))

    doc.build(story)
    return buf.getvalue()


def to_docx_bytes(report: PreMortemReport) -> bytes:
    import docx
    from docx.shared import Pt, RGBColor

    document = docx.Document()
    document.add_heading("PreMortem AI - Procurement Risk Report", level=0)

    table = document.add_table(rows=0, cols=2)
    table.style = "Light Grid Accent 1"
    rows = [
        ("Procurement", report.procurement_name),
        ("Equipment", report.equipment_type),
        ("Contract Value", f"Rs {report.contract_value_cr:.2f} Cr"),
        ("Overall Risk Score", f"{report.overall_risk_score:.0f}/100"),
        ("Failure Probability", f"{report.failure_probability_pct:.0f}%"),
        ("Confidence", f"{report.confidence_pct:.0f}%"),
        ("Predicted Delay", f"{report.predicted_delay_months:.0f} months"),
        ("Projected Loss", f"Rs {report.projected_financial_loss_cr:.2f} Cr"),
    ]
    for k, v in rows:
        cells = table.add_row().cells
        cells[0].text = k
        cells[1].text = str(v)

    p = document.add_paragraph()
    run = p.add_run(f"Recommended Decision: {report.recommended_decision.value}")
    run.bold = True
    run.font.size = Pt(14)
    color = {
        "GO": RGBColor(0x16, 0xA3, 0x4A),
        "GO WITH CONDITIONS": RGBColor(0xD9, 0x77, 0x06),
        "NO-GO": RGBColor(0xDC, 0x26, 0x26),
    }.get(report.recommended_decision.value, RGBColor(0xDC, 0x26, 0x26))
    run.font.color.rgb = color

    document.add_heading("Predicted Failure Mode", level=1)
    document.add_paragraph(report.predicted_failure_mode)

    document.add_heading("Supporting Evidence", level=1)
    for e in report.supporting_evidence:
        document.add_paragraph(str(e), style="List Bullet")

    document.add_heading("Predicted Outcomes", level=1)
    for o in report.predicted_outcomes:
        document.add_paragraph(str(o), style="List Bullet")

    if report.conditions:
        document.add_heading("Conditions for Approval", level=1)
        for c in report.conditions:
            document.add_paragraph(str(c), style="List Bullet")

    document.add_heading("Agent Findings", level=1)
    for r in report.agent_results:
        document.add_heading(
            f"{r.agent} - Risk {r.risk_score:.0f}/100 ({r.risk_level.value})",
            level=2,
        )
        for f in r.findings:
            document.add_paragraph(str(f), style="List Bullet")
        if r.reasoning:
            document.add_paragraph(r.reasoning)

    buf = io.BytesIO()
    document.save(buf)
    return buf.getvalue()
